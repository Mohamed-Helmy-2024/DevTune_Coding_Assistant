import os
from typing import List, Dict
from pathlib import Path
import logging
import os
import hashlib

try:
    # Use langchain loaders where available
    from langchain.document_loaders import TextLoader, PyPDFLoader, Docx2txtLoader
except Exception:
    # Fallback if langchain isn't available at import time (unit tests, etc.)
    TextLoader = None
    PyPDFLoader = None
    Docx2txtLoader = None


class Document:
    """Class to represent a document with its content and metadata"""
    def __init__(self, content: str, metadata: Dict = None):
        self.content = content
        self.metadata = metadata or {}


class DocumentLoader:
    """Base class for loading documents from various file formats"""

    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.supported_formats = ['.txt', '.pdf', '.docx', '.md']

    def load_document(self, file_path: str) -> Document:
        """Load a single document from file path"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_ext = Path(file_path).suffix.lower()

        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_ext}")

        if file_ext == '.txt' or file_ext == '.md':
            return self._load_text_file(file_path)
        elif file_ext == '.pdf':
            return self._load_pdf_file(file_path)
        elif file_ext == '.docx':
            return self._load_docx_file(file_path)

    def _load_text_file(self, file_path: str) -> Document:
        """Load text file"""
        try:
            # Use langchain TextLoader if available, which handles unicode and encodings
            if TextLoader is not None:
                loader = TextLoader(file_path, encoding='utf-8')
                docs = loader.load()
                content = "\n".join([d.page_content for d in docs])
            else:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()

            metadata = {
                'source': file_path,
                'file_type': 'text',
                'file_name': Path(file_path).name
            }
            # Compute simple hash for deduplication
            metadata['doc_hash'] = self._compute_hash(content)

            return Document(content=content, metadata=metadata)
        except Exception as e:
            self.logger.error(f"Error loading text file {file_path}: {e}")
            raise

    def _load_pdf_file(self, file_path: str) -> Document:
        """Load PDF file"""
        try:
            # Prefer LangChain's PyPDFLoader if available
            if PyPDFLoader is not None:
                loader = PyPDFLoader(file_path)
                docs = loader.load()
                content = "\n".join([d.page_content for d in docs])
                num_pages = len(docs)
            else:
                import PyPDF2
                content = ""
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    num_pages = len(pdf_reader.pages)
                    for page_num in range(num_pages):
                        page = pdf_reader.pages[page_num]
                        content += page.extract_text() + "\n"

            metadata = {
                'source': file_path,
                'file_type': 'pdf',
                'file_name': Path(file_path).name,
                'num_pages': num_pages
            }
            metadata['doc_hash'] = self._compute_hash(content)

            return Document(content=content, metadata=metadata)
        except ImportError:
            self.logger.error("PyPDF2 not installed. Install it with: pip install PyPDF2")
            raise
        except Exception as e:
            self.logger.error(f"Error loading PDF file {file_path}: {e}")
            raise

    def _load_docx_file(self, file_path: str) -> Document:
        """Load DOCX file"""
        try:
            if Docx2txtLoader is not None:
                loader = Docx2txtLoader(file_path)
                docs = loader.load()
                content = "\n".join([d.page_content for d in docs])
                # We don't get paragraphs/table counts from docx loader easily; set 0 as fallback
                num_paragraphs = 0
                num_tables = 0
            else:
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                content_parts = []
                # Extract paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content_parts.append(paragraph.text)
                # Extract tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            content_parts.append(" | ".join(row_text))
                content = "\n".join(content_parts)
                num_paragraphs = len(doc.paragraphs)
                num_tables = len(doc.tables)

            metadata = {
                'source': file_path,
                'file_type': 'docx',
                'file_name': Path(file_path).name,
                'num_paragraphs': num_paragraphs,
                'num_tables': num_tables
            }
            metadata['doc_hash'] = self._compute_hash(content)

            return Document(content=content, metadata=metadata)
        except ImportError:
            self.logger.error("python-docx not installed. Install it with: pip install python-docx")
            raise
        except Exception as e:
            self.logger.error(f"Error loading DOCX file {file_path}: {e}")
            raise

    def load_directory(self, directory_path: str) -> List[Document]:
        """Load all supported documents from a directory"""
        if not os.path.exists(directory_path):
            raise FileNotFoundError(f"Directory not found: {directory_path}")

        documents = []
        for file_name in os.listdir(directory_path):
            file_path = os.path.join(directory_path, file_name)

            if os.path.isfile(file_path):
                file_ext = Path(file_path).suffix.lower()
                if file_ext in self.supported_formats:
                    try:
                        doc = self.load_document(file_path)
                        documents.append(doc)
                    except Exception as e:
                        self.logger.warning(f"Skipping file {file_path}: {e}")

        return documents

    def _compute_hash(self, content: str) -> str:
        try:
            h = hashlib.sha256(content.encode('utf-8')).hexdigest()
            return h
        except Exception:
            return ""
